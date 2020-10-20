import speech_recognition as sr
import os
from pydub import AudioSegment
from pydub.utils import make_chunks
from docx import Document
import moviepy.editor as mp
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def audio_to_word(path):
    myaudio = AudioSegment.from_file(fr"{ path}", "wav")
    chunk_length_ms = 80000  # pydub calculates in millisecond
    chunks = make_chunks(myaudio, chunk_length_ms)  # Make chunks of 80 sec

    # Export all of the individual  chunks as wav files

    for i, chunk in enumerate(chunks):
        chunk_name = "chunk{0}.wav".format(i)
        print("exporting", chunk_name)
        chunk.export(chunk_name, format="wav")
    #create word file
    document = Document()
    document.add_heading('a', 0)
    i = 0
    for chunk in chunks:
        print("in chunks")
        filename = 'chunk' + str(i) + '.wav'
        print("Processing chunk " + str(i))
        file = filename
        r = sr.Recognizer()
        with sr.AudioFile(file) as source:
            # r.adjust_for_ambie nt_noise(source)
            audio_listened = r.record(source)
        try:
            rec = r.recognize_google(audio_listened, language="iw-IL")
            p = document.add_paragraph(rec)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.alignment = 3

        except:
            rec = r.recognize_google(audio_listened, show_all=True)
            print(rec, type(rec))
        os.remove(file)
        i += 1
    os.remove("converted.wav")
    document.save('recognize.docx')


def video_to_wav(path):
    clip = mp.VideoFileClip(fr"{path}")
    return clip.audio.write_audiofile(r"converted.wav")
def name_of_file(path):
    s = path
    s1= s.split('/')[-1:][0]
    return s1

def to_text(path):
    video_to_wav(path)
    audio_to_word("converted.wav")

name_of_file("sgsgd/sdgsdg/ss")



